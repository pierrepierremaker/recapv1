from io import BytesIO
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter


def _build_meta_block(meta: dict | None) -> list[str]:
    """
    Construit une petite liste de lignes texte à partir des métadonnées.
    """
    if not meta:
        return []

    lines = []

    titre = meta.get("title")
    date = meta.get("date")
    lieu = meta.get("location")
    participants = meta.get("participants")

    if titre:
        lines.append(f"Titre : {titre}")
    if date:
        lines.append(f"Date : {date}")
    if lieu:
        lines.append(f"Lieu : {lieu}")
    if participants:
        lines.append(f"Participants : {participants}")

    return lines


def export_docx(text: str, meta: dict | None = None) -> BytesIO:
    """
    Génère un DOCX avec un en-tête contenant les métadonnées de réunion
    puis le corps du compte rendu.
    """
    buffer = BytesIO()
    doc = Document()

    doc.add_heading("Compte rendu de réunion", level=1)

    # Bloc métadonnées
    meta_lines = _build_meta_block(meta)
    if meta_lines:
        for line in meta_lines:
            doc.add_paragraph(line)
        doc.add_paragraph("")  # ligne vide de séparation

    # Corps du CR
    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")

    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_pdf(text: str, meta: dict | None = None) -> BytesIO:
    """
    Génère un PDF avec un en-tête contenant les métadonnées de réunion
    puis le corps du compte rendu.
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()

    story = []

    # Titre
    story.append(Paragraph("Compte rendu de réunion", styles["Title"]))
    story.append(Spacer(1, 12))

    # Bloc métadonnées
    meta_lines = _build_meta_block(meta)
    if meta_lines:
        for line in meta_lines:
            story.append(Paragraph(line, styles["BodyText"]))
        story.append(Spacer(1, 12))

    # Corps du CR
    for line in text.split("\n"):
        if line.strip():
            story.append(Paragraph(line, styles["BodyText"]))
        else:
            story.append(Spacer(1, 6))

    doc.build(story)
    buffer.seek(0)
    return buffer
